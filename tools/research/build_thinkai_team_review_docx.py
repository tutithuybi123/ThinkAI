"""Build the Vietnamese ThinkAI v1.2 team-review DOCX from its Markdown source."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "proposals" / "ThinkAI-Idea-Team-Review-v1.2.md"
OUTPUT = ROOT / "docs" / "proposals" / "ThinkAI-Idea-Team-Review-v1.2.docx"

NAVY = "16324F"
BLUE = "2367A2"
TEAL = "148A8A"
PALE_BLUE = "EAF3F8"
PALE_TEAL = "E7F4F2"
PALE_GOLD = "FFF4D8"
LIGHT = "F4F6F8"
MID = "D6DEE5"
TEXT = "24313D"
MUTED = "5E6C78"
WHITE = "FFFFFF"

# Keep only the two major transitions explicit. Other sections flow naturally so
# the review stays compact and tables are not stranded by decorative page breaks.
# Keep the Capability Receipt and Kill/Reframe tables with their headings rather
# than leaving a repeated table header orphaned at a page boundary.
PAGE_BREAK_SECTIONS: set[str] = {"7.", "21."}


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)")


def add_inline(paragraph, text: str, size: float | None = None, color: str = TEXT) -> None:
    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.color.rgb = RGBColor.from_string(color)
            if size:
                run.font.size = Pt(size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
            if size:
                run.font.size = Pt(size)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor.from_string(BLUE)
            if size:
                run.font.size = Pt(size - 0.5)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            clean = token.rstrip(".,);]")
            suffix = token[len(clean):]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                paragraph.add_run(suffix)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 18, NAVY, 12, 7),
        ("Heading 2", 13, BLUE, 9, 4),
        ("Heading 3", 11, TEAL, 7, 3),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        styles[list_style].font.name = "Aptos"
        styles[list_style]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        styles[list_style].font.size = Pt(9)
        styles[list_style].paragraph_format.space_after = Pt(2)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "THINKAI v1.2  |  TÀI LIỆU ĐÁNH GIÁ Ý TƯỞNG NỘI BỘ"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("ThinkAI v1.2 • 14/08/2026   |   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(56)
    p.add_run("THINKAI").font.color.rgb = RGBColor.from_string(TEAL)
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(title, "ĐỀ XUẤT & PHẢN BIỆN Ý TƯỞNG", size=30)
    subtitle = doc.add_paragraph(style="Subtitle")
    add_inline(subtitle, "Tài liệu để nhóm và giáo viên cùng quyết định hướng phát triển", size=14, color=MUTED)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    cell = band.cell(0, 0)
    set_cell_width(cell, 17.0)
    shade(cell, PALE_TEAL)
    set_cell_margins(cell, 260, 280, 260, 280)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("SỨ MỆNH")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    para2 = cell.add_paragraph()
    add_inline(para2, "ThinkAI phải giúp người dùng thực sự học được điều gì đó, thay vì chỉ giúp họ lấy được đáp án nhanh hơn.", size=15, color=NAVY)
    para2.runs[0].bold = True

    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(quote, "“Học nó. Dùng nó ở chỗ khác. Biến nó thành kỹ năng của mình.”", size=18, color=BLUE)
    quote.runs[0].italic = True

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    rows = [
        ("Mục đích", "Đánh giá có tiếp tục phát triển ThinkAI hay không"),
        ("Đối tượng đọc", "Thành viên nhóm và giáo viên hướng dẫn"),
        ("Trạng thái", "Nghiên cứu / đề xuất — chưa phải hồ sơ dự thi"),
        ("Ngày", "14/08/2026 • v1.2"),
    ]
    for i, (left, right) in enumerate(rows):
        set_cell_width(meta.rows[i].cells[0], 4.0)
        set_cell_width(meta.rows[i].cells[1], 12.8)
        shade(meta.rows[i].cells[0], NAVY)
        shade(meta.rows[i].cells[1], LIGHT if i % 2 == 0 else WHITE)
        for cell in meta.rows[i].cells:
            set_cell_margins(cell, 90, 130, 90, 130)
        p1 = meta.rows[i].cells[0].paragraphs[0]
        p1.text = left
        p1.runs[0].font.color.rgb = RGBColor.from_string(WHITE)
        p1.runs[0].bold = True
        p2 = meta.rows[i].cells[1].paragraphs[0]
        p2.text = right

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(35)
    add_inline(note, "Tài liệu này phân biệt rõ: CÓ CƠ SỞ NGHIÊN CỨU / ĐÃ THỬ NỘI BỘ / ĐỀ XUẤT / TƯƠNG LAI.", size=9, color=MUTED)
    doc.add_page_break()


def add_status_legend(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    entries = [
        ("CÓ CƠ SỞ", PALE_TEAL, TEAL),
        ("ĐÃ THỬ NỘI BỘ", PALE_BLUE, BLUE),
        ("ĐỀ XUẤT", PALE_GOLD, "A36A00"),
        ("TƯƠNG LAI", LIGHT, MUTED),
    ]
    for idx, (label, fill, color) in enumerate(entries):
        cell = table.cell(0, idx)
        set_cell_width(cell, 4.15)
        shade(cell, fill)
        set_cell_margins(cell, 100, 70, 100, 70)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [p.strip() for p in lines[idx].strip().strip("|").split("|")]
        rows.append(parts)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", p.replace(" ", "")) for p in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    available = 17.2
    widths = [available / col_count] * col_count
    if col_count == 2:
        widths = [available * 0.31, available * 0.69]
    elif col_count == 3:
        widths = [available * 0.25, available * 0.25, available * 0.50]
    elif col_count == 4:
        widths = [available * 0.20, available * 0.25, available * 0.25, available * 0.30]
    for r_idx, values in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, NAVY if r_idx == 0 else (LIGHT if r_idx % 2 == 0 else WHITE))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            text = values[c_idx] if c_idx < len(values) else ""
            add_inline(p, text, size=8.2, color=WHITE if r_idx == 0 else TEXT)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
        table.rows[r_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 16.8)
    set_cell_margins(cell, 150, 200, 150, 200)
    shade(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size=11.5, color=NAVY)
    for run in p.runs:
        run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_numbered_item(doc: Document, text: str) -> None:
    # Keep the source number literal. Word's built-in List Number style otherwise
    # continues numbering across unrelated sections (including the bibliography).
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.48)
    p.paragraph_format.first_line_indent = Cm(-0.48)
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text, size=9)


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_status_legend(doc)

    # Skip source title/front matter and begin at section 1.
    start = next(i for i, line in enumerate(lines) if line.startswith("# 1."))
    idx = start
    first_h1 = True
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line or line == "---":
            idx += 1
            continue
        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            prefix = title.split(maxsplit=1)[0] if title else ""
            should_break = not first_h1 and (prefix in PAGE_BREAK_SECTIONS or title.startswith("Nguồn"))
            if should_break:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, title, size=18, color=NAVY)
            first_h1 = False
            idx += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[3:].strip(), size=13, color=BLUE)
            idx += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[4:].strip(), size=11, color=TEAL)
            idx += 1
            continue
        if line.startswith("> "):
            quote_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("> "):
                quote_lines.append(lines[idx].strip()[2:])
                idx += 1
            add_callout(doc, " ".join(quote_lines))
            continue
        if line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip(), size=9)
            idx += 1
            continue
        if re.match(r"^\d+\.\s", line):
            add_numbered_item(doc, line)
            idx += 1
            continue
        p = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**"):
            p.paragraph_format.space_before = Pt(2)
        add_inline(p, line, size=9)
        idx += 1

    core = doc.core_properties
    core.title = "ThinkAI v1.2 — Tài liệu đề xuất và phản biện ý tưởng dành cho nhóm"
    core.subject = "Controlled learning evidence / AI-assisted learning verification"
    core.author = "ThinkAI team — AI-assisted research synthesis"
    core.keywords = "ThinkAI, capability receipt, assistance exposure, learning evidence events"
    core.comments = "Internal review document; not a competition submission."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
